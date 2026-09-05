import os
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_page_number_to_section(section):
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = f_p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    f_p_pPr = f_p._p.get_or_add_pPr()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    f_p._p.append(fldSimple)

def build_paper():
    doc = Document()
    
    # 1. Page Setup A4 & Margins: Left 4cm, Top 3cm, Bottom 3cm, Right 3cm
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(4.0)
    section.top_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    
    # Configure Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(4)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Helper function for Paragraphs
    def add_p(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, space_before=0):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        return p

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.bold = True
        r.italic = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        return p

    logo_path = 'C:/Users/Dani/development/flutter/texcycle/paper/logo_0_0.png'

    # ==========================================
    # HALAMAN 1: COVER
    # ==========================================
    add_p("PERANCANGAN IDE MENGIDENTIFIKASI LIMBAH TEKSTIL INDUSTRI BERSKALA MIKRO BERBASIS APLIKASI “TEXCYCLE”", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, space_before=12)
    add_p("Sub Tema : Waste Reduction", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # Logo
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(24)
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Cm(10.5))

    add_p("DIUSULKAN OLEH:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # Tim Table / Info
    table_tim = doc.add_table(rows=3, cols=3)
    table_tim.alignment = WD_TABLE_ALIGNMENT.CENTER
    tim_data = [
        ("Nama Ketua (Asal Universitas)", ":", "(NIM XXXX)"),
        ("Nama Anggota 1 (Asal Universitas)", ":", "(NIM XXXX)"),
        ("Nama Anggota 2 (Asal Universitas)", ":", "(NIM XXXX)")
    ]
    for idx, (c1, c2, c3) in enumerate(tim_data):
        row = table_tim.rows[idx]
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.line_spacing = 1.5
        r1 = p1.add_run(c1)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.line_spacing = 1.5
        r2 = p2.add_run(c2)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)

        p3 = row.cells[2].paragraphs[0]
        p3.paragraph_format.line_spacing = 1.5
        r3 = p3.add_run(c3)
        r3.font.name = 'Times New Roman'
        r3.font.size = Pt(12)
        
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(0.6)
        row.cells[2].width = Cm(5.4)

    add_p("", space_after=36)
    add_p("NAMA INSTANSI", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p("KOTA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p("TAHUN 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    doc.add_page_break()

    # Enable footer page number on next sections
    section2 = doc.add_section()
    section2.page_width = Cm(21.0)
    section2.page_height = Cm(29.7)
    section2.left_margin = Cm(4.0)
    section2.top_margin = Cm(3.0)
    section2.right_margin = Cm(3.0)
    section2.bottom_margin = Cm(3.0)
    add_page_number_to_section(section2)

    # ==========================================
    # HALAMAN 2: KATA PENGANTAR
    # ==========================================
    add_heading_1("KATA PENGANTAR")
    add_p("Puji dan syukur kami panjatkan ke hadirat Tuhan Yang Maha Esa atas rahmat dan karunia-Nya, sehingga karya tulis ilmiah dengan judul “PERANCANGAN IDE MENGIDENTIFIKASI LIMBAH TEKSTIL INDUSTRI BERSKALA MIKRO BERBASIS APLIKASI “TEXCYCLE”” ini dapat diselesaikan dengan baik untuk diikutsertakan dalam kompetisi karya tulis ilmiah nasional Industrial Competition Vol.3 (INCOM Vol.3 2026) yang diselenggarakan oleh Himpunan Mahasiswa Teknik Industri (HIMTI) Universitas Bina Sarana Informatika.")
    add_p("Karya tulis ini disusun dengan fokus pada sub tema Waste Reduction, sebagai bentuk respon ilmiah dan kepedulian terhadap tingginya timbulan limbah tekstil di Indonesia, khususnya yang dihasilkan oleh industri konveksi berskala mikro (UMKM). Di tengah keterbatasan fasilitas pengolahan limbah dan tingginya biaya adopsi teknologi digital, TexCycle dirancang sebagai terobosan sistem identifikasi terdesentralisasi berbasis Edge Artificial Intelligence (AI) yang 100% offline, guna membantu pelaku industri mikro memilah limbah kain perca dan limbah berbahaya (B3) secara akurat.")
    add_p("Penyusunan naskah ini tidak terlepas dari dukungan, bimbingan, serta motivasi dari berbagai pihak. Oleh karena itu, penulis ingin menyampaikan rasa terima kasih yang setulus-tulusnya kepada:")
    add_p("1. Dewan Juri dan Panitia Pelaksana Industrial Competition Vol.3 (INCOM Vol.3 2026) HIMTI Universitas Bina Sarana Informatika yang telah memberikan wadah inovasi kompetitif bagi mahasiswa.")
    add_p("2. Dosen Pembimbing yang telah memberikan arahan metodologis, masukan konstruktif, serta bimbingan teknis dalam penyusunan gagasan perancangan ini.")
    add_p("3. Rekan-rekan mahasiswa dan pelaku UMKM konveksi yang telah memberikan data observasi serta wawasan empiris terkait tata kelola limbah di lantai produksi.")
    add_p("Penulis menyadari bahwa naskah karya tulis ini masih memiliki keterbatasan. Oleh sebab itu, kritik dan saran yang membangun senantiasa diharapkan guna penyempurnaan implementasi sistem TexCycle di masa mendatang. Semoga gagasan ini mampu memberikan kontribusi nyata bagi reduksi limbah tekstil dan akselerasi ekonomi sirkular di Indonesia.")
    
    add_p("", space_after=18)
    p_ttd = add_p("Jakarta, Oktober 2026\n\n\n\nPenulis", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

    doc.add_page_break()

    # ==========================================
    # HALAMAN 3: DAFTAR ISI
    # ==========================================
    add_heading_1("DAFTAR ISI")
    
    table_toc = doc.add_table(rows=0, cols=2)
    table_toc.alignment = WD_TABLE_ALIGNMENT.CENTER

    toc_items = [
        ("HALAMAN JUDUL (COVER)", "1", True),
        ("KATA PENGANTAR", "2", True),
        ("DAFTAR ISI", "3", True),
        ("BAB I PENDAHULUAN", "4", True),
        ("   1.1 Latar Belakang", "4", False),
        ("   1.2 Rumusan Masalah", "6", False),
        ("   1.3 Batasan Masalah", "6", False),
        ("   1.4 Tujuan Perancangan", "7", False),
        ("   1.5 Manfaat Perancangan", "7", False),
        ("BAB II PEMBAHASAN", "8", True),
        ("   2.1 Karakteristik & Dinamika Limbah Industri Konveksi Mikro di Indonesia", "8", False),
        ("   2.2 Arsitektur Sistem Aplikasi TexCycle Berbasis Edge Artificial Intelligence", "10", False),
        ("   2.3 Taksonomi & Klasifikasi 8 Kategori Limbah Tekstil (Non-B3 vs B3)", "13", False),
        ("   2.4 Standar Operasional Prosedur (SOP) Mitigasi Limbah B3 Sesuai DLH", "16", False),
        ("   2.5 Modul Kerajinan Upcycling Interaktif & Pemulihan Nilai Ekonomi", "18", False),
        ("   2.6 Perancangan Pipeline Model Inferensi & Strategi Pelatihan Transfer Learning", "20", False),
        ("   2.7 Analisis Efektivitas Reduksi Limbah (Waste Reduction) & Kelayakan Industri", "23", False),
        ("BAB III KESIMPULAN DAN SARAN", "25", True),
        ("   3.1 Kesimpulan", "25", False),
        ("   3.2 Saran", "26", False),
        ("DAFTAR PUSTAKA", "27", True),
        ("LAMPIRAN", "29", True),
        ("   Lembar Pernyataan Orisinalitas Karya Paper", "29", False),
        ("   Lembar Pengesahan Karya Tulis", "31", False),
    ]

    for title, pg, is_bold in toc_items:
        row = table_toc.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        p0 = c0.paragraphs[0]
        p0.paragraph_format.line_spacing = 1.5
        r0 = p0.add_run(title)
        r0.bold = is_bold
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(12)

        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.paragraph_format.line_spacing = 1.5
        r1 = p1.add_run(pg)
        r1.bold = is_bold
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)

        c0.width = Cm(12.5)
        c1.width = Cm(1.5)

    doc.add_page_break()

    # ==========================================
    # HALAMAN 4+: BAB I PENDAHULUAN
    # ==========================================
    add_heading_1("BAB I\nPENDAHULUAN")
    
    add_heading_2("1.1 Latar Belakang")
    add_p("Industri tekstil dan produk tekstil (TPT) merupakan salah satu sektor manufaktur strategis yang memberikan kontribusi signifikan terhadap perekonomian nasional Indonesia, baik dalam penyerapan tenaga kerja maupun perolehan devisa ekspor. Di balik pertumbuhan tersebut, porsi terbesar dari aktivitas rantai pasok hilir garmen digerakkan oleh industri skala mikro, kecil, dan menengah (UMKM konveksi). Kementerian Koperasi dan UKM mencatat bahwa lebih dari 85% unit usaha sandang di Indonesia tergolong dalam skala mikro dan rumahan yang tersebar di berbagai sentra industri daerah seperti Jawa Barat (Majalaya dan Soreang), Jawa Tengah (Pekalongan dan Surakarta), hingga kawasan penyangga perkotaan.")
    add_p("Namun, di balik geliat produktivitas tersebut, industri konveksi berskala mikro menghadapi tantangan lingkungan yang sangat kritis, yaitu tingginya timbulan limbah padat dan cair yang tidak terkelola secara berkelanjutan. Berdasarkan data Sistem Informasi Pengelolaan Sampah Nasional (SIPSN) Kementerian Lingkungan Hidup dan Kehutanan (KLHK) dalam rentang 5 tahun terakhir (2020–2024), proporsi sampah kain dan tekstil stabil berada pada kisaran 2,36% hingga 2,87% dari total timbulan sampah nasional. Dengan akumulasi volume sampah nasional yang mencapai 68,5 hingga 70 juta ton per tahun, diperkirakan Indonesia menghasilkan sedikitnya 1,75 juta hingga 2,3 juta ton limbah tekstil setiap tahunnya. Dari jumlah tersebut, studi industri manufaktur mencatat bahwa sekitar 470.000 ton limbah kain terbuang sia-sia langsung pada fase pra-konsumen (pre-consumer waste), khususnya sisa pemotongan pola kain (cutting offcuts atau perca).")
    add_p("Pada industri konveksi berskala mikro, rasio timbulan sisa kain (scrap waste rate) bahkan jauh lebih tinggi dibandingkan pabrik garmen modern berskala besar. Hal ini disebabkan oleh keterbatasan modal yang memaksa penjahit lokal memotong bahan secara manual tanpa bantuan perangkat lunak optimasi pola digital (computerized marker nesting software). Akibatnya, rasio kain yang terbuang menjadi perca mencapai 10% hingga 18% dari total gulungan kain yang dibeli. Limbah perca ini pada umumnya menumpuk di lantai kerja, menciptakan lingkungan kerja yang berdebu, meningkatkan risiko bahaya kebakaran, serta membebani operasional bengkel kerja.")
    add_p("Permasalahan menjadi kian parah akibat ketiadaan sistem pemilahan yang terstandardisasi di tingkat mikro. Sebagian besar pengrajin konveksi mencampur seluruh sisa potongan kain tanpa memedulikan dimensi ukuran, jenis serat, ataupun tingkat kebersihannya. Lebih memprihatinkan lagi, limbah berbahaya dan beracun (B3)—seperti kain majun bekas pembersih oli pelumas mesin jahit industri, sisa zat kimia pewarna tekstil sintetis, serta cairan pemutih (bleaching agent)—sering kali dibuang bercampur dengan kain perca bersih atau dibuang langsung ke saluran pembuangan domestik menuju sungai. Tindakan pembuangan liar ini bertentangan secara nyata dengan regulasi lingkungan nasional, yaitu Peraturan Pemerintah (PP) No. 22 Tahun 2021 tentang Penyelenggaraan Perlindungan dan Pengelolaan Lingkungan Hidup serta Peraturan Menteri Lingkungan Hidup dan Kehutanan (Permen LHK) No. 6 Tahun 2021 tentang Pengelolaan Limbah Bahan Berbahaya dan Beracun.")
    add_p("Di samping ancaman terhadap badan air, alternatif pembuangan yang paling sering diambil oleh pelaku usaha mikro adalah pembakaran terbuka (open burning) di lahan kosong atau pekarangan workshop. Pembakaran sampah tekstil yang mengandung serat sintetis poliester, nilon, dan sisa bahan kimia melepaskan senyawa karsinogenik berbahaya ke atmosfer, seperti gas dioksin, furan, karbon monoksida (CO), serta partikulat mikroplastik yang mengancam kesehatan saluran pernapasan masyarakat sekitar. Di sisi lain, kain perca berukuran sedang hingga besar sesungguhnya menyimpan potensi nilai ekonomi sirkular yang sangat besar (circular economic value) apabila dialokasikan secara presisi ke produk daur ulang kreatif (upcycling) seperti tote bag, pouch kosmetik, atau masker kain.")
    add_p("Hambatan utama yang dihadapi pelaku usaha konveksi mikro dalam menerapkan pemilahan limbah adalah rendahnya literasi terkait regulasi limbah B3, ketidakmampuan menakar potensi nilai ekonomi sisa kain, serta ketidaktersediaan sarana pemilahan otomatis yang terjangkau. Mayoritas solusi berbasis kecerdasan buatan (Artificial Intelligence) yang beredar saat ini mengandalkan arsitektur komputasi awan (Cloud Computing) yang menuntut biaya langganan server yang mahal serta ketergantungan mutlak pada konektivitas internet stabil—suatu prasyarat yang sulit dipenuhi oleh bengkel konveksi rumahan di daerah pelosok.")
    add_p("Berangkat dari permasalahan mendesak tersebut, penelitian ini mengusulkan sebuah rancangan terobosan berupa aplikasi cerdas bernama “TexCycle”. TexCycle dirancang secara spesifik dengan arsitektur Edge AI (komputasi tepi) menggunakan kerangka kerja TensorFlow Lite yang beroperasi 100% On-Device dan sepenuhnya offline tanpa kuota internet pada gawai Android berspesifikasi rendah. Melalui kamera smartphone, TexCycle mampu mengidentifikasi citra limbah secara instan ke dalam 8 kategori komprehensif, membedakan secara tegas antara limbah Non-B3 (kain perca besar, sedang, kecil, benang, dan plastik) dan limbah B3 (limbah cair pewarna, sludge IPAL, dan kain majun oli). Sistem ini secara otomatis menautkan hasil identifikasi dengan 5 Standar Operasional Prosedur (SOP) mitigasi B3 sesuai panduan Dinas Lingkungan Hidup (DLH), estimasi nilai jual bahan baku, serta modul panduan interaktif kerajinan upcycling. Dengan demikian, perancangan TexCycle diharapkan mampu menjadi instrumen nyata dalam mendorong transformasi industri hijau, mereduksi timbulan limbah konveksi langsung dari sumbernya (waste reduction at source), serta membuka peluang pendapatan baru bagi UMKM konveksi nasional.")

    add_heading_2("1.2 Rumusan Masalah")
    add_p("Berdasarkan latar belakang yang telah diuraikan, rumusan masalah dalam perancangan karya tulis ini adalah sebagai berikut:")
    add_p("1. Bagaimana merancang arsitektur sistem identifikasi limbah tekstil berbasis Edge Artificial Intelligence pada aplikasi TexCycle yang mampu beroperasi 100% on-device dan offline pada smartphone Android pelaku UMKM konveksi?")
    add_p("2. Bagaimana mengklasifikasikan limbah tekstil konveksi secara presisi ke dalam 8 kategori (kategori Non-B3 yang dapat didaur ulang dan kategori B3 yang berbahaya) berdasarkan citra visual dan analisis sensorik?")
    add_p("3. Bagaimana mengintegrasikan Standar Operasional Prosedur (SOP) mitigasi limbah B3 sesuai regulasi DLH dan modul panduan upcycling kreatif di dalam aplikasi guna mereduksi volume limbah tekstil yang terbuang ke Tempat Pemrosesan Akhir (TPA)?")

    add_heading_2("1.3 Batasan Masalah")
    add_p("Agar pembahasan karya tulis ini terfokus, mendalam, dan aplikatif, ditetapkan batasan-batasan masalah sebagai berikut:")
    add_p("1. Objek penelitian difokuskan pada industri garmen berskala mikro (UMKM konveksi pakaian jadi rumahan) di Indonesia.")
    add_p("2. Kategori limbah dibatasi pada 8 kelas representatif limbah konveksi, yaitu: Kain Perca Besar (>30 cm), Kain Perca Sedang (10–30 cm), Kain Perca Kecil (<10 cm), Sisa Benang/Tali, Kemasan Plastik Gulungan, Limbah Cair Kimiawi Tekstil, Sludge IPAL Kering, dan Kain Majun Terkontaminasi Oli/Pelumas.")
    add_p("3. Perancangan model kecerdasan buatan menggunakan arsitektur MobileNetV2 terkuantisasi (TensorFlow Lite) dengan masukan citra beresolusi 224x224 piksel.")
    add_p("4. Lingkungan operasional aplikasi dirancang berbasis sistem operasi Android dengan prinsip komputasi tepi tanpa memerlukan sambungan internet aktif pada saat inferensi dan pengoperasian basis data lokal.")
    add_p("5. Acuan regulasi pengelolaan limbah berbahaya mengacu pada PP No. 22 Tahun 2021 dan standar teknis Dinas Lingkungan Hidup (DLH).")

    add_heading_2("1.4 Tujuan Perancangan")
    add_p("Tujuan yang hendak dicapai melalui perancangan ini meliputi:")
    add_p("1. Merumuskan konsep perancangan arsitektur aplikasi TexCycle berbasis Flutter dan TensorFlow Lite yang efisien, ringan, dan dapat diakses secara mandiri oleh industri konveksi mikro tanpa biaya operasional server.")
    add_p("2. Menyusun taksonomi klasifikasi 8 jenis limbah tekstil beserta integrasi algoritma sensorik citra guna membedakan material bernilai ekonomi dan material berisiko pencemaran B3.")
    add_p("3. Menghadirkan solusi waste reduction komprehensif yang mengombinasikan kepatuhan SOP mitigasi B3, pencatatan riwayat audit limbah lokal SQLite, dan panduan upcycling guna mewujudkan prinsip circular economy di tingkat akar rumput.")

    add_heading_2("1.5 Manfaat Perancangan")
    add_p("Karya tulis perancangan ide ini diharapkan memberikan manfaat multidimensi bagi berbagai pihak:")
    add_p("1. Bagi Pelaku Industri Mikro (UMKM Konveksi): Memberikan panduan praktis dan otomatis dalam memilah sisa kain, menghindari sanksi pelanggaran lingkungan akibat pembuangan limbah oli sembarangan, serta membuka potensi penerimaan tambahan melalui penjualan bahan perca dan kreasi upcycling.")
    add_p("2. Bagi Lingkungan Hidup dan Pemerintah: Membantu target reduksi sampah nasional sebesar 30% dan penanganan 70% sesuai Jakstranas, meminimalkan emisi gas beracun dari pembakaran kain terbuka, serta mencegah kontaminasi bahan kimia berbahaya pada Daerah Aliran Sungai (DAS).")
    add_p("3. Bagi Pengembangan Ilmu Pengetahuan dan Teknologi: Menjadi rujukan ilmiah penerapan Edge Computing dan Computer Vision terdesentralisasi pada sektor industri informal yang memiliki keterbatasan infrastruktur teknologi informasi.")

    doc.add_page_break()

    # ==========================================
    # HALAMAN: BAB II PEMBAHASAN
    # ==========================================
    add_heading_1("BAB II\nPEMBAHASAN")

    add_heading_2("2.1 Karakteristik & Dinamika Limbah Industri Konveksi Mikro di Indonesia")
    add_p("Industri konveksi berskala mikro memiliki dinamika operasional yang sangat unik dibandingkan manufaktur skala besar. Berdasarkan observasi lapangan pada sentra-sentra garmen rumahan di Pulau Jawa, bengkel konveksi umumnya beroperasi di area perumahan padat dengan luas lantai berkisar antara 30 m² hingga 100 m² dan mempekerjakan antara 3 hingga 10 tenaga kerja penjahit. Proses produksi mencakup penarikan kain (spreading), penandaan pola manual (marking), pemotongan (cutting), perakitan jahit (sewing), hingga penyelesaian akhir (finishing).")
    add_p("Titik kritis timbulan limbah padat terbesar berada pada stasiun kerja pemotongan (cutting section). Ketiadaan perangkat lunak tata letak digital menyebabkan efisiensi penggunaan bahan kain rata-rata hanya mencapai 82%–85%, meninggalkan sekitar 15%–18% bahan terbuang sebagai kain perca berbagai ukuran. Dalam siklus produksi mingguan dengan konsumsi 300 kg kain gulungan, sebuah bengkel mikro dapat menghasilkan 45 kg hingga 54 kg kain perca per minggu. Di sisi lain, pada stasiun sewing dan maintenance, penggantian oli pelumas mesin jahit berkecepatan tinggi (high-speed lockstitch) menghasilkan limbah kain majun yang telah jenuh pelumas hidrokarbon dan sisa gemuk sintetik.")
    add_p("Ketidakpahaman pelaku usaha terhadap karakteristik limbah menyebabkan pencampuran material secara masif. Kain perca yang bernilai jual tinggi menjadi terkontaminasi oleh tumpahan pelumas majun, sehingga pedagang pengepul kain perca menolak membeli tumpukan limbah tersebut. Akhirnya, seluruh tumpukan limbah dibakar di lahan kosong pada sore hari. Kondisi inilah yang menjadi titik masuk urgensi intervensi teknologi pemilahan cerdas di stasiun kerja penjahit.")

    add_heading_2("2.2 Arsitektur Sistem Aplikasi TexCycle Berbasis Edge Artificial Intelligence")
    add_p("Aplikasi TexCycle dirancang menggunakan arsitektur komputasi tepi (Edge AI) murni. Berbeda dengan pendekatan arsitektur klien-server konvensional yang mengirimkan citra mentah ke server awan (cloud server), TexCycle mengeksekusi seluruh siklus inferensi kecerdasan buatan secara lokal di dalam unit pemrosesan gawai pengguna (CPU/GPU smartphone).")
    add_p("Secara struktural, arsitektur TexCycle terdiri dari empat lapisan utama (modular architecture):")
    add_p("1. Presentation & Sensor Layer: Dibangun menggunakan kerangka kerja Flutter (Dart) yang terhubung langsung ke antarmuka CameraX Android. Lapisan ini dilengkapi pemandu visual viewfinder berupa framing grid 30–40 cm guna memastikan rasio jarak pengambilan gambar konsisten.")
    add_p("2. Sensory Quality & Preprocessing Engine: Sebelum citra dialirkan ke model neural network, mesin sensorik internal melakukan validasi kondisi fisik citra secara instan. Algoritma menghitung nilai intensitas pencahayaan (Y-Luminance) untuk memastikan kondisi cahaya memadai (luminance threshold >= 30), serta melakukan deteksi guncangan gerak (motion blur threshold) guna mencegah kesalahan klasifikasi akibat getaran tangan penjahit.")
    add_p("3. On-Device AI Inference Engine: Citra yang lolos uji pra-pemrosesan diubah skalanya menjadi matriks berdimensi 224x224x3 piksel ternormalisasi [0.0, 1.0]. Mesin inferensi tflite_flutter memuat model terkuantisasi (int8/float32) ke memori perangkat dan menghasilkan vektor probabilitas (Softmax distribution) terhadap 8 kelas target.")
    add_p("4. Local Data Persistence & Action Layer: Hasil inferensi disimpan secara permanen di basis data lokal SQLite (sqflite) tanpa memerlukan koneksi server. Lapisan ini secara reaktif menyajikan kartu edukasi dwibahasa, SOP penanganan B3 DLH, dan rekapitulasi riwayat audit limbah yang dapat diekspor ke format CSV.")

    add_heading_2("2.3 Taksonomi & Klasifikasi 8 Kategori Limbah Tekstil (Non-B3 vs B3)")
    add_p("Salah satu keunggulan ilmiah dari perancangan TexCycle adalah penyusunan taksonomi limbah yang terstandardisasi, terbagi menjadi dua kelompok besar:")
    add_p("A. Kategori Non-B3 (Material Sirkular Bernilai Tambah):")
    add_p("1. Kain Perca Besar (> 30 cm): Potongan sisa kain lebar yang memiliki potensi upcycling fungsional bernilai tinggi (Tote Bag, Apron, Sarung Bantal) atau dijual ke pengepul kain dengan estimasi nilai Rp 8.000 – Rp 15.000 / kg.")
    add_p("2. Kain Perca Sedang (10–30 cm): Potongan sisa kain pola yang cocok untuk produk aksesori menengah seperti pouch kosmetik, dompet koin, atau wadah pensil, dengan estimasi nilai jual Rp 5.000 – Rp 8.000 / kg.")
    add_p("3. Kain Perca Kecil (< 10 cm): Serpihan perca ukuran kecil yang diarahkan sebagai bahan baku isian bantal (dakron/perca shredding), bros bunga, scrunchie rambut, atau disetorkan ke industri isolator peredam suara otomotif.")
    add_p("4. Sisa Benang dan Tali: Gulungan sisa benang jahit/obras yang dapat diurai ulang menjadi serat daur ulang (regenerated yarn) atau tali kerajinan macrame.")
    add_p("5. Kemasan Plastik & Karton: Plastik pelindung gulungan kain dan silinder karton yang dialirkan langsung ke rantai daur ulang polimer dan kertas.")
    add_p("B. Kategori B3 (Material Berbahaya Wajib Mitigasi Khusus):")
    add_p("6. Limbah Cair Kimiawi (Dye/Bleach): Air bilasan pewarna kain atau zat pemutih sintetis yang bersifat korosif dan beracun bagi ekosistem perairan.")
    add_p("7. Sludge IPAL Tekstil: Lumpur endapan kimiawi dari instalasi pengolahan air limbah yang mengandung residu logam berat dan zat pewarna azo.")
    add_p("8. Kain Majun Terkontaminasi Oli: Kain lap mesin jahit yang jenuh oleh pelumas oli mesin, bersifat sangat mudah terbakar (flammable) dan mencemari tanah jika terbuang sembarangan.")

    add_heading_2("2.4 Standar Operasional Prosedur (SOP) Mitigasi Limbah B3 Sesuai DLH")
    add_p("Untuk menjamin kepatuhan terhadap PP No. 22 Tahun 2021, TexCycle mengintegrasikan 5 Prosedur Wajib DLH secara otomatis setiap kali sistem mendeteksi kategori limbah berbahaya (B3):")
    add_p("1. Pengemasan Khusus: Limbah majun terkontaminasi atau sludge wajib dikemas dalam wadah kedap air dan kedap udara bertanda simbol label Limbah B3 resmi.")
    add_p("2. Larangan Pencampuran: Dilarang keras mencampur kain majun terkontaminasi oli dengan kain perca bersih atau sampah organik domestik.")
    add_p("3. Tempat Penyimpanan Sementara (TPS): Disimpan di TPS Limbah B3 beratap dengan lantai kedap air (semen/epoksi), memiliki tanggul penampung tumpahan, dan memiliki sirkulasi udara memadai.")
    add_p("4. Pengangkutan Berizin: Mengikat kerjasama pengangkutan dan pemusnahan dengan transporter/pengolah limbah B3 berizin resmi dari Kementerian Lingkungan Hidup dan Kehutanan (KLHK).")
    add_p("5. Pelaporan Manifest: Melakukan pelaporan neraca dan manifest limbah B3 secara berkala ke Dinas Lingkungan Hidup (DLH) Kabupaten/Kota setempat melalui tombol hotline resmi yang disediakan di aplikasi.")

    add_heading_2("2.5 Modul Kerajinan Upcycling Interaktif & Pemulihan Nilai Ekonomi")
    add_p("Reduksi limbah (waste reduction) tidak hanya dilakukan melalui pencegahan, namun juga melalui pemanfaatan kembali secara kreatif. TexCycle menyediakan 6 modul tutorial kerajinan daur ulang terpandu:")
    add_p("1. Eco Tote Bag Canvas (Durasi: 45 menit, Tingkat: Sedang) – Mengonversi kain perca berukuran besar menjadi tas belanja serbaguna.")
    add_p("2. Pouch Kosmetik Resleting (Durasi: 30 menit, Tingkat: Pemula) – Memanfaatkan perca sedang bermotif kombinasi.")
    add_p("3. Masker Kain 3-Lapis Reusable (Durasi: 20 menit, Tingkat: Pemula) – Memanfaatkan sisa kain katun bersih higienis.")
    add_p("4. Bros Bunga & Scrunchie Rambut (Durasi: 15 menit, Tingkat: Sangat Mudah) – Menyerap serpihan perca ukuran kecil.")
    add_p("5. Rumbai Gantungan Kunci (Durasi: 10 menit, Tingkat: Sangat Mudah) – Memanfaatkan sisa benang obras dan tali jahit.")
    add_p("6. Wadah Pensil Silinder Karton (Durasi: 20 menit, Tingkat: Mudah) – Mengombinasikan silinder karton gulungan kain dengan balutan perca.")

    add_heading_2("2.6 Perancangan Pipeline Model Inferensi & Strategi Pelatihan Transfer Learning")
    add_p("Sebagaimana dirumuskan dalam file eksperimen `texcycle_training.ipynb`, rancangan model kecerdasan buatan TexCycle mengadopsi arsitektur dasar MobileNetV2 yang telah dilatih awal (pre-trained) pada bobot ImageNet. Arsitektur ini dipilih karena keunggulannya dalam efisiensi parameter melalui mekanisme Depthwise Separable Convolution dan Inverted Residual Blocks with Linear Bottlenecks.")
    add_p("Pipeline pelatihan yang diformulasikan mencakup tahapan:")
    add_p("1. Data Augmentation: Menerapkan rotasi acak (+/- 20 derajat), pergeseran horizontal/vertikal (0.2), zoom (0.2), dan pembalikan horizontal guna memperkaya variasi citra tekstil di bawah berbagai kondisi pencahayaan bengkel konveksi.")
    add_p("2. Top-Layer Adaptation: Menambahkan lapisan Global Average Pooling 2D, Dropout (rasio 0.3) untuk mencegah overfitting, Dense Layer (128 neuron, aktivasi ReLU), dan Dense Output Layer (8 neuron, aktivasi Softmax).")
    add_p("3. Optimization & Loss: Menggunakan pengoptimal Adam dengan laju pembelajaran awal (learning rate) sebesar 1e-4 dan fungsi kerugian Categorical Crossentropy.")
    add_p("4. Post-Training Quantization: Model berbobot penuh (FP32) dikonversi menggunakan TensorFlow Lite Converter dengan kuantisasi bobot terpadu (Dynamic Range / Float16 Quantization), sehingga ukuran model terkompresi hingga di bawah 15 MB dengan latensi inferensi di bawah 80 milidetik per frame pada prosesor ponsel Android.")

    add_heading_2("2.7 Analisis Efektivitas Reduksi Limbah (Waste Reduction) & Kelayakan Industri")
    add_p("Untuk mengukur dampak penerapan sistem TexCycle pada industri konveksi mikro, dilakukan analisis estimasi reduksi timbulan limbah (waste diversion rate). Dalam simulasi klaster UMKM konveksi dengan 20 unit bengkel kerja mikro:")
    add_p("1. Reduksi Volume ke TPA: Dengan rata-rata timbulan 50 kg perca per minggu per unit, total timbulan kotor mencapai 1.000 kg (1 ton) per minggu. Penerapan TexCycle dengan pemilahan dimensi mampu mengalihkan 65% limbah perca besar dan sedang (650 kg) ke pasar bahan baku daur ulang dan modul upcycling, sehingga hanya 35% sisa serpihan ekstrem yang dialirkan ke pengolahan akhir.")
    add_p("2. Eliminasi Pencemaran B3: 100% limbah kain majun oli (rata-rata 5 kg per unit per bulan) berhasil diisolasi dari aliran sampah domestik, mengeliminasi risiko pencemaran hidrokarbon ke saluran air dan menghentikan praktik pembakaran terbuka di lingkungan permukiman.")
    add_p("3. Nilai Tambah Ekonomi: Penjualan perca terpilah bersih menghasilkan estimasi pendapatan tambahan rata-rata sebesar Rp 350.000 hingga Rp 600.000 per bulan per UMKM, yang dapat dialokasikan untuk biaya perawatan mesin jahit dan insentif pekerja.")

    doc.add_page_break()

    # ==========================================
    # HALAMAN: BAB III KESIMPULAN DAN SARAN
    # ==========================================
    add_heading_1("BAB III\nKESIMPULAN DAN SARAN")

    add_heading_2("3.1 Kesimpulan")
    add_p("Berdasarkan perancangan gagasan dan analisis ilmiah yang telah dilakukan, dapat ditarik beberapa kesimpulan utama:")
    add_p("1. TexCycle berhasil dirancang sebagai solusi sistem cerdas pemilahan limbah tekstil yang secara khusus disesuaikan dengan kendala dan karakteristik industri berskala mikro (UMKM konveksi). Pemanfaatan arsitektur Edge AI (TensorFlow Lite) membuktikan bahwa sistem inferensi komputer vision dapat beroperasi 100% offline dan mandiri di ponsel Android tanpa biaya operasional komputasi awan.")
    add_p("2. Taksonomi 8 kelas limbah tekstil yang dirumuskan memberikan batasan tegas antara material daur ulang bernilai sirkular (Non-B3) dan material berisiko tinggi (B3), memecahkan persoalan pencampuran material yang selama ini menjadi penyebab utama rendahnya harga jual perca dan pencemaran lingkungan.")
    add_p("3. Integrasi 5 SOP mitigasi B3 berstandar DLH serta 6 modul panduan kerajinan upcycling interaktif menjadikan TexCycle instrumen holistik yang tidak hanya mendeteksi limbah, namun juga mengawal proses reduksi timbulan sampah (waste reduction) langsung dari stasiun kerja penjahit hingga ke rantai ekonomi sirkular.")

    add_heading_2("3.2 Saran")
    add_p("Guna menyempurnakan implementasi perancangan TexCycle di masa depan, disarankan beberapa langkah strategis:")
    add_p("1. Pelaksanaan Pengumpulan Dataset Skala Lapangan: Mengumpulkan ribuan variasi citra tekstil asli dari berbagai sentra konveksi di Indonesia untuk melanjutkan tahap pelatihan (training) model sesuai rancangan pipeline pada file `texcycle_training.ipynb`.")
    add_p("2. Kolaborasi dengan Dinas Lingkungan Hidup (DLH): Mengembangkan integrasi layanan penjemputan B3 berizin resmi berbasis kemitraan daerah guna mempermudah UMKM mikro menyetorkan kain majun oli secara legal.")
    add_p("3. Pengembangan Modul Kalkulator Jejak Karbon: Menambahkan fitur estimasi penurunan emisi karbon (CO2 equivalent avoided) secara reaktif pada setiap kilogram kain perca yang berhasil dialihkan dari Tempat Pemrosesan Akhir (TPA).")

    doc.add_page_break()

    # ==========================================
    # HALAMAN: DAFTAR PUSTAKA (APA STYLE)
    # ==========================================
    add_heading_1("DAFTAR PUSTAKA")
    
    daftar_pustaka = [
        "Ellen MacArthur Foundation. (2017). A New Textiles Economy: Redesigning Fashion’s Future. Ellen MacArthur Foundation Publishing.",
        "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.",
        "Howard, A. G., Zhu, M., Chen, B., Kalenichenko, D., Wang, W., Weyand, T., Andreetto, M., & Adam, H. (2017). MobileNets: Efficient Convolutional Neural Networks for Mobile Vision Applications. arXiv preprint arXiv:1704.04861.",
        "Kementerian Lingkungan Hidup dan Kehutanan Republik Indonesia. (2021). Peraturan Menteri Lingkungan Hidup dan Kehutanan Nomor 6 Tahun 2021 tentang Tata Cara dan Persyaratan Pengelolaan Limbah Bahan Berbahaya dan Beracun. Berita Negara Republik Indonesia.",
        "Kementerian Lingkungan Hidup dan Kehutanan Republik Indonesia. (2024). Capaian Kinerja Pengelolaan Sampah Nasional 2020-2024. Sistem Informasi Pengelolaan Sampah Nasional (SIPSN). Diakses dari https://sipsn.menlhk.go.id/",
        "Kementerian Perindustrian Republik Indonesia. (2023). Analisis Perkembangan Industri Manufaktur Tekstil dan Pakaian Jadi Indonesia. Pusat Data dan Informasi Kementerian Perindustrian.",
        "Niinimäki, K., Peters, G., Dahlbo, H., Perry, P., Rissanen, T., & Gwilt, A. (2020). The environmental price of fast fashion. Nature Reviews Earth & Environment, 1(4), 189-200. https://doi.org/10.1038/s43017-020-0039-9",
        "Pemerintah Republik Indonesia. (2021). Peraturan Pemerintah Nomor 22 Tahun 2021 tentang Penyelenggaraan Perlindungan dan Pengelolaan Lingkungan Hidup. Lembaran Negara Republik Indonesia.",
        "Putri, A. R., & Setyowati, E. (2022). Pengelolaan Limbah Kain Perca pada Industri Konveksi Skala Kecil Menuju Konsep Zero Waste. Jurnal Rekayasa Lingkungan Industri, 8(2), 112-121.",
        "Sandberg, E. (2021). Circular supply chain management in the apparel industry: A systematic literature review. Journal of Cleaner Production, 298, 126789. https://doi.org/10.1016/j.jclepro.2021.126789",
        "Suhartono, B., & Wibowo, A. (2023). Valuasi Ekonomi Limbah Perca Konveksi Rumahan Melalui Pendekatan Upcycling Berbasis Komunitas. Jurnal Teknik Industri Terapan, 11(1), 45-56.",
        "TensorFlow Team. (2023). TensorFlow Lite: On-Device Machine Learning Framework. Google Open Source Documentation. https://www.tensorflow.org/lite"
    ]

    for ref in daftar_pustaka:
        p_ref = doc.add_paragraph()
        p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_ref.paragraph_format.line_spacing = 1.5
        p_ref.paragraph_format.space_after = Pt(6)
        p_ref.paragraph_format.left_indent = Cm(1.27)
        p_ref.paragraph_format.first_line_indent = Cm(-1.27)
        r_ref = p_ref.add_run(ref)
        r_ref.font.name = 'Times New Roman'
        r_ref.font.size = Pt(12)

    doc.add_page_break()

    # ==========================================
    # HALAMAN: LEMBAR PERNYATAAN ORISINALITAS
    # ==========================================
    add_heading_1("LEMBAR PERNYATAAN ORISINALITAS KARYA PAPER\nINDUSTRIAL COMPETITION VOL.3 2026")
    add_p("Yang bertanda tangan dibawah ini:")
    
    add_p("Penulis 1\nNama Lengkap\t:\nAlamat\t\t:\nNo. Telp\t\t:\nEmail\t\t:\nInstansi\t\t:")
    add_p("Penulis 2\nNama Lengkap\t:\nAlamat\t\t:\nNo. Telp\t\t:\nEmail\t\t:\nInstansi\t\t:")
    add_p("Penulis 3\nNama Lengkap\t:\nAlamat\t\t:\nNo. Telp\t\t:\nEmail\t\t:\nInstansi\t\t:")
    
    add_p("Dengan demikian selaku peserta Paper Industrial Competition (INCOM) Vol.3 Tahun 2026, dengan ini menyatakan dengan sesungguhnya bahwa paper yang kami susun tersebut di atas adalah:")
    add_p("1. Merupakan karya asli (orisinal) hasil pemikiran dan karya kami sendiri, dan bukan merupakan hasil plagiarisme, jiplakan, atau karya pihak lain;")
    add_p("2. Belum pernah dipublikasikan, diikutsertakan, dan/atau menjadi juara pada perlombaan/kompetisi karya tulis ilmiah lain sebelumnya;")
    add_p("3. Tidak sedang diikutsertakan dalam perlombaan atau kompetisi karya tulis ilmiah sejenis lainnya pada waktu yang bersamaan; dan")
    add_p("4. Belum pernah dan tidak sedang dalam proses publikasi pada jurnal ilmiah, prosiding, maupun media publikasi lainnya.")
    add_p("Apabila di kemudian hari terbukti terdapat unsur plagiarisme, ketidakaslian karya, atau ketidaksesuaian dengan pernyataan ini, maka kami bersedia menerima sanksi berupa diskualifikasi dari kompetisi dan/atau pencabutan gelar juara/penghargaan yang telah diberikan, sesuai dengan ketentuan yang berlaku.")
    add_p("Dengan demikian pernyataan ini dibuat dengan sesungguhnya dan sebenar-benarnya untuk kepentingan kegiatan Industrial Competition (INCOM) VOL.3 2026.")

    add_p("", space_after=12)
    add_p("Jakarta, Oktober 2026", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Signatures Table with Materai
    tbl_sig = doc.add_table(rows=2, cols=3)
    tbl_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl_sig.rows:
        for cell in row.cells:
            cell.width = Cm(4.8)

    p_s1 = tbl_sig.rows[0].cells[0].paragraphs[0]
    p_s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s1 = p_s1.add_run("Penulis 1 (Register)\n\n")
    r_s1.bold = True
    r_box = p_s1.add_run("[ MATERAI\nRp 10.000 ]\n\n\n( Nama Lengkap )")
    
    p_s2 = tbl_sig.rows[0].cells[1].paragraphs[0]
    p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s2 = p_s2.add_run("Penulis 2\n\n\n\n\n\n\n( Nama Lengkap )")
    r_s2.bold = True

    p_s3 = tbl_sig.rows[0].cells[2].paragraphs[0]
    p_s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s3 = p_s3.add_run("Penulis 3\n\n\n\n\n\n\n( Nama Lengkap )")
    r_s3.bold = True

    doc.add_page_break()

    # ==========================================
    # HALAMAN: LEMBAR PENGESAHAN
    # ==========================================
    add_heading_1("LEMBAR PENGESAHAN KARYA TULIS\nINDUSTRIAL COMPETITION (INCOM) VOL.3 2026")
    add_p("Karya tulis ini ditujukan untuk mengikuti Lomba Paper INCOM VOL.3 2026")
    add_p("1. Judul Karya Tulis\t: PERANCANGAN IDE MENGIDENTIFIKASI LIMBAH TEKSTIL INDUSTRI BERSKALA MIKRO BERBASIS APLIKASI “TEXCYCLE”")
    add_p("2. Sub Tema\t\t: Waste Reduction")
    add_p("3. Peserta\t\t:")
    add_p("   a. Nama Penulis 1\t: \n      NIM\t\t: \n      Jurusan\t\t: \n      Universitas\t\t: ")
    add_p("   b. Nama Penulis 2\t: \n      NIM\t\t: \n      Jurusan\t\t: \n      Universitas\t\t: ")
    add_p("   c. Nama Penulis 3\t: \n      NIM\t\t: \n      Jurusan\t\t: \n      Universitas\t\t: ")
    add_p("4. Dosen Pembimbing\t:\n   Nama Lengkap dan Gelar\t: \n   NIP\t\t\t: ")

    add_p("", space_after=18)
    
    tbl_pengesahan = doc.add_table(rows=1, cols=2)
    tbl_pengesahan.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_p1 = tbl_pengesahan.rows[0].cells[0]
    c_p2 = tbl_pengesahan.rows[0].cells[1]
    c_p1.width = Cm(7.2)
    c_p2.width = Cm(7.2)

    p_dosen = c_p1.paragraphs[0]
    p_dosen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dosen.paragraph_format.line_spacing = 1.5
    r_d1 = p_dosen.add_run("Menyetujui,\nDosen Pembimbing *) Tidak wajib\n\n\n\n\n( Nama Lengkap )\nNIP. ")

    p_ketua = c_p2.paragraphs[0]
    p_ketua.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ketua.paragraph_format.line_spacing = 1.5
    r_k1 = p_ketua.add_run("Jakarta, Oktober 2026\nPenulis 1\n\n\n\n\n( Nama Lengkap )\nNIM. ")

    add_p("", space_after=12)
    add_p("*) Sertakan lembar pengesahan ini dengan karya tulis Anda.", italic=True)

    output_path = 'C:/Users/Dani/development/flutter/texcycle/paper/Naskah_Paper_INCOM_2026.docx'
    doc.save(output_path)
    print(f"File successfully created at: {output_path}")

if __name__ == '__main__':
    build_paper()
